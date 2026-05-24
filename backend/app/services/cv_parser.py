"""CV parser service (Phase 11).

Extracts plain text from PDF / DOCX (and optionally image-based CVs via
Tesseract OCR when installed), then applies a set of conservative
heuristic extractors to pull out structured candidate data:

- Name, email, mobile
- Nationality, current location
- Total experience (years), GCC / Qatar experience
- Expected salary, notice period, visa status
- Skills, education entries, certifications, languages
- Previous companies / employers
- The full extracted text (stored for later AI review in Phase 13)

The parser is intentionally rule-based — no LLM call, no network — so
it is safe to run synchronously inside the upload request. Phase 13
will layer an AI review on top of the data we extract here.

All extractors are deliberately permissive: when in doubt they return
``None`` (or an empty list) so that HR can fill the field manually
without being misled by a wrong guess.
"""
from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence


PARSER_VERSION = "phase11.heuristics.1"


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


class CvParseError(Exception):
    """Raised when text extraction fails for a known file type."""


def extract_text(path: str | Path, extension: Optional[str] = None) -> str:
    """Return the plain-text contents of a CV file.

    ``extension`` should be the lowercase suffix (no leading dot). If
    omitted, it is inferred from the file's path.
    """
    p = Path(path)
    ext = (extension or p.suffix.lstrip(".")).lower()
    if not p.exists():
        raise CvParseError(f"CV file not found: {p}")

    if ext == "pdf":
        return _extract_pdf(p)
    if ext == "docx":
        return _extract_docx(p)
    if ext in {"png", "jpg", "jpeg"}:
        return _extract_image_ocr(p)
    if ext == "doc":
        # Legacy .doc binary format is not supported by python-docx and
        # would need antiword/textract. Surface this clearly so HR
        # knows to re-save as DOCX.
        raise CvParseError(
            "Legacy .doc files are not supported — please save the CV as "
            ".docx or .pdf and re-upload."
        )
    raise CvParseError(f"Unsupported CV extension: {ext!r}")


def _extract_pdf(p: Path) -> str:
    try:
        from pypdf import PdfReader  # imported lazily
    except ImportError as exc:  # pragma: no cover
        raise CvParseError("pypdf is not installed") from exc

    try:
        reader = PdfReader(str(p))
    except Exception as exc:  # noqa: BLE001
        raise CvParseError(f"Could not open PDF: {exc}") from exc

    chunks: List[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — some pages have no text layer
            chunks.append("")
    return _normalize_whitespace("\n".join(chunks))


def _extract_docx(p: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise CvParseError("python-docx is not installed") from exc

    try:
        document = docx.Document(str(p))
    except Exception as exc:  # noqa: BLE001
        raise CvParseError(f"Could not open DOCX: {exc}") from exc

    lines: List[str] = [para.text for para in document.paragraphs]
    # Also walk simple tables (common in CVs).
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return _normalize_whitespace("\n".join(lines))


def _extract_image_ocr(p: Path) -> str:
    """Run Tesseract OCR on an image-based CV.

    Returns an empty string (rather than raising) when either Pillow or
    pytesseract is missing, so the rest of the pipeline can still store
    the document — HR can fill the fields manually.
    """
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except ImportError:
        return ""

    try:
        with Image.open(p) as img:
            text = pytesseract.image_to_string(img)
    except Exception:  # noqa: BLE001 — Tesseract binary missing, bad image, etc.
        return ""
    return _normalize_whitespace(text)


def _normalize_whitespace(text: str) -> str:
    # Collapse repeated whitespace within a line but preserve line breaks.
    out_lines = []
    for line in text.splitlines():
        line = unicodedata.normalize("NFKC", line)
        line = re.sub(r"[ \t ]+", " ", line).strip()
        out_lines.append(line)
    # Drop runs of blank lines.
    deduped: List[str] = []
    prev_blank = False
    for line in out_lines:
        if not line:
            if prev_blank:
                continue
            prev_blank = True
        else:
            prev_blank = False
        deduped.append(line)
    return "\n".join(deduped).strip()


# ---------------------------------------------------------------------------
# Heuristic field extractors
# ---------------------------------------------------------------------------


EMAIL_RE = re.compile(r"\b[\w.+-]+@[\w-]+(?:\.[\w-]+)+\b")
# Phone numbers across GCC / international formats — at least 7 digits,
# optional leading + and country code, spaces / dashes / dots allowed.
PHONE_RE = re.compile(
    r"(?:(?<!\d)(?:\+?\d{1,3}[\s.\-]?)?(?:\(?\d{2,4}\)?[\s.\-]?)?\d{3,4}[\s.\-]?\d{3,5})"
)
URL_RE = re.compile(r"https?://\S+|www\.\S+", re.IGNORECASE)

# Tokens that signal an "experience" mention.
EXP_RE = re.compile(
    r"(\d{1,2}(?:\.\d)?)\s*\+?\s*(?:years?|yrs?)\b",
    re.IGNORECASE,
)

# Salary / compensation — capture numeric amount + currency hint.
SALARY_RE = re.compile(
    r"(?:salary|expected|expectation|ctc|package)[^\n]{0,40}?"
    r"(?:(?:qar|qr|aed|sar|inr|usd|gbp|eur)\s*)?"
    r"(\d{1,3}(?:[,\s]\d{3})+|\d{4,7})"
    r"(?:\s*(qar|qr|aed|sar|inr|usd|gbp|eur|/month|per month|pm))?",
    re.IGNORECASE,
)

NOTICE_RE = re.compile(
    r"notice\s*period[^\n]{0,40}?"
    r"(immediate(?:ly)?|"
    r"(?:[0-9]+|one|two|three|four|six|seven|eight)\s*(?:day|week|month)s?)",
    re.IGNORECASE,
)

VISA_RE = re.compile(
    r"\bvisa(?:\s*status)?\b[^\n]{0,80}|"
    r"\b(?:transferable\s*(?:noc|rp)|nor?c\s+available|nor?c\s+required|"
    r"work\s*permit|residen(?:t|ce)\s*permit|rp\s*holder|own\s*visa|"
    r"family\s*visa|sponsorship\s*required|visit\s*visa|tourist\s*visa)\b",
    re.IGNORECASE,
)

DEGREE_TOKENS = [
    "PhD", "Ph.D", "Doctorate",
    "Master of", "Master's", "Masters", "MBA", "MSc", "M.Sc", "MA", "M.A",
    "M.Tech", "MEng", "M.Eng", "MS ",
    "Bachelor of", "Bachelor's", "Bachelors", "BSc", "B.Sc",
    "BA", "B.A", "B.Tech", "BEng", "B.Eng", "BBA", "BCA", "B.Com",
    "Diploma", "Higher National Diploma", "HND",
    "Higher Secondary", "Secondary School", "High School",
    "Associate Degree",
]

CERT_TOKENS = [
    "PMP", "PRINCE2", "ITIL", "Six Sigma", "Lean Six Sigma",
    "CISSP", "CISA", "CISM", "CEH", "AWS Certified", "Azure Certified",
    "GCP Certified", "Google Cloud Certified", "OCJP", "OCP", "MCSA", "MCSE",
    "CFA", "CPA", "ACCA", "CMA", "SHRM", "CIPD", "PHR", "SPHR",
    "Scrum Master", "PSM", "CSM", "Scrum Product Owner",
    "TOGAF", "Salesforce Certified", "Cisco Certified", "CCNA", "CCNP",
    "Oracle Certified", "Red Hat Certified", "RHCSA", "RHCE",
]

SKILLS_DICT = [
    # Programming languages
    "Python", "Java", "JavaScript", "TypeScript", "C#", "C++", "C ", "Go",
    "Rust", "Ruby", "PHP", "Kotlin", "Swift", "Objective-C", "Scala",
    "R ", "MATLAB", "Perl", "Bash", "Shell",
    # Web / frameworks
    "React", "Next.js", "Angular", "Vue", "Svelte", "Node.js", "Express",
    "Django", "Flask", "FastAPI", "Spring Boot", "Spring", "ASP.NET",
    ".NET", "Laravel", "Rails", "GraphQL", "REST", "gRPC", "HTML", "CSS",
    "SASS", "Tailwind",
    # Data / ML
    "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch",
    "Kafka", "RabbitMQ", "Snowflake", "BigQuery", "Hadoop", "Spark",
    "Airflow", "ETL", "Power BI", "Tableau", "Excel",
    "Pandas", "NumPy", "TensorFlow", "PyTorch", "scikit-learn",
    "Machine Learning", "Deep Learning", "NLP",
    # Cloud / DevOps
    "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform",
    "Ansible", "Jenkins", "GitLab CI", "GitHub Actions", "Linux",
    "Nginx", "Apache",
    # Business / retail / FMCG (relevant for PUG roles)
    "Retail Operations", "Merchandising", "Visual Merchandising",
    "Category Management", "Sales", "Distribution", "FMCG",
    "Logistics", "Supply Chain", "Inventory Management", "Procurement",
    "Vendor Management", "Customer Service", "Cashiering", "POS",
    "Forecasting", "Demand Planning",
    # Engineering / construction
    "AutoCAD", "Revit", "MEP", "HVAC", "Plumbing", "Electrical",
    "Project Management", "Construction Management", "Estimation",
    "BOQ", "Civil Engineering", "Structural Engineering",
    # Automotive
    "Mechanical Engineering", "Diesel Engines", "Hydraulic Systems",
    "Turbochargers", "Diagnostics", "Body Work",
    # Finance / HR / Real estate
    "Accounting", "Financial Reporting", "Tax", "Audit", "Budgeting",
    "Payroll", "Recruitment", "Talent Acquisition", "Onboarding",
    "Employee Relations", "Property Sales", "Leasing", "Brokerage",
]

LANGUAGE_DICT = [
    "Arabic", "English", "Hindi", "Urdu", "Malayalam", "Tamil",
    "Telugu", "Kannada", "Bengali", "Sinhala", "Tagalog", "Filipino",
    "Nepali", "Pashto", "French", "Spanish", "German", "Russian",
    "Mandarin", "Chinese", "Japanese", "Korean", "Italian", "Portuguese",
    "Turkish", "Persian", "Farsi", "Punjabi", "Marathi",
]

# Common nationalities seen in the GCC labour market.
NATIONALITY_DICT = [
    "Indian", "Filipino", "Filipina", "Pakistani", "Bangladeshi",
    "Sri Lankan", "Nepali", "Nepalese", "Egyptian", "Jordanian",
    "Lebanese", "Syrian", "Sudanese", "Tunisian", "Moroccan",
    "Algerian", "Saudi", "Qatari", "Emirati", "Bahraini", "Omani",
    "Kuwaiti", "Yemeni", "Iranian", "Iraqi", "Turkish", "Afghan",
    "British", "American", "Canadian", "Australian", "South African",
    "Kenyan", "Ugandan", "Ethiopian", "Nigerian", "Ghanaian",
    "Indonesian", "Malaysian", "Singaporean", "Thai", "Vietnamese",
]

GULF_CITIES = [
    "Doha", "Qatar", "Lusail", "Al Rayyan", "Al Wakrah", "Mesaieed",
    "Dubai", "Abu Dhabi", "Sharjah", "Ajman", "Ras Al Khaimah",
    "UAE", "Saudi Arabia", "Riyadh", "Jeddah", "Dammam", "Khobar",
    "Manama", "Bahrain", "Kuwait", "Kuwait City",
    "Muscat", "Oman", "Salalah",
]

# Loose section headers used to find the "experience" / "education" parts
# of a CV without depending on exact wording.
SECTION_HEADERS = {
    "experience": (
        "experience", "work experience", "professional experience",
        "employment history", "career history", "work history",
    ),
    "education": (
        "education", "qualifications", "academic", "academic background",
        "academic qualifications", "educational qualifications",
    ),
    "skills": (
        "skills", "technical skills", "core skills", "key skills",
        "competencies", "core competencies",
    ),
    "certifications": (
        "certifications", "certificates", "professional certifications",
        "trainings", "courses",
    ),
    "languages": ("languages", "languages known", "language proficiency"),
}


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass(slots=True)
class ParsedEducation:
    raw: str
    degree: Optional[str] = None
    institution: Optional[str] = None
    year: Optional[int] = None


@dataclass(slots=True)
class ParsedCompany:
    name: str
    title: Optional[str] = None
    duration: Optional[str] = None


@dataclass(slots=True)
class ParsedCv:
    """Bundle returned by `parse_cv` / `parse_text`."""

    full_text: str
    parser_version: str = PARSER_VERSION

    # Identity
    name: Optional[str] = None
    email: Optional[str] = None
    mobile: Optional[str] = None

    # Demographics
    nationality: Optional[str] = None
    current_location: Optional[str] = None

    # Career snapshot
    current_designation: Optional[str] = None
    current_company: Optional[str] = None
    total_experience_years: Optional[float] = None
    gcc_experience_years: Optional[float] = None
    qatar_experience_years: Optional[float] = None

    # Compensation / availability
    expected_salary: Optional[int] = None
    notice_period: Optional[str] = None
    visa_status: Optional[str] = None

    # Structured lists
    skills: List[str] = field(default_factory=list)
    education: List[ParsedEducation] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    languages: List[str] = field(default_factory=list)
    previous_companies: List[ParsedCompany] = field(default_factory=list)

    def education_as_json(self) -> List[Dict[str, Any]]:
        return [
            {
                "raw": e.raw,
                "degree": e.degree,
                "institution": e.institution,
                "year": e.year,
            }
            for e in self.education
        ]

    def companies_as_json(self) -> List[Dict[str, Any]]:
        return [
            {"name": c.name, "title": c.title, "duration": c.duration}
            for c in self.previous_companies
        ]


# ---------------------------------------------------------------------------
# Top-level orchestrators
# ---------------------------------------------------------------------------


def parse_cv(file_path: str | Path, extension: Optional[str] = None) -> ParsedCv:
    """Extract + parse a CV from disk."""
    text = extract_text(file_path, extension=extension)
    return parse_text(text)


def parse_text(text: str) -> ParsedCv:
    """Parse a pre-extracted plain-text CV."""
    if not text:
        return ParsedCv(full_text="")

    lines = [line for line in text.splitlines() if line.strip()]
    lower = text.lower()

    parsed = ParsedCv(full_text=text)
    parsed.email = _first_match(EMAIL_RE, text)
    parsed.mobile = _extract_mobile(text)
    parsed.name = _extract_name(lines, email=parsed.email)
    parsed.nationality = _extract_one_of(text, NATIONALITY_DICT)
    parsed.current_location = _extract_location(lines)

    parsed.total_experience_years = _extract_years(EXP_RE, text)
    parsed.gcc_experience_years = _extract_regional_years(
        text, region_terms=("gcc", "gulf", "middle east")
    )
    parsed.qatar_experience_years = _extract_regional_years(
        text, region_terms=("qatar", "doha")
    )

    parsed.expected_salary = _extract_salary(text)
    parsed.notice_period = _extract_notice(text)
    parsed.visa_status = _extract_visa(text)

    sections = _split_sections(lines)
    parsed.skills = _extract_skills(text, sections.get("skills"))
    parsed.languages = _extract_languages(text, sections.get("languages"))
    parsed.certifications = _extract_certifications(text, sections.get("certifications"))
    parsed.education = _extract_education(sections.get("education", []))
    parsed.previous_companies = _extract_companies(sections.get("experience", []))

    if parsed.previous_companies:
        most_recent = parsed.previous_companies[0]
        parsed.current_company = parsed.current_company or most_recent.name
        parsed.current_designation = parsed.current_designation or most_recent.title

    return parsed


# ---------------------------------------------------------------------------
# Field-specific helpers
# ---------------------------------------------------------------------------


def _first_match(pattern: re.Pattern[str], text: str) -> Optional[str]:
    m = pattern.search(text)
    return m.group(0).strip() if m else None


def _extract_mobile(text: str) -> Optional[str]:
    """Pick the first phone-shaped run with at least 7 digits."""
    # Skip the local part of email addresses (they contain digits too).
    cleaned = EMAIL_RE.sub(" ", text)
    cleaned = URL_RE.sub(" ", cleaned)
    for raw in PHONE_RE.findall(cleaned):
        digits = re.sub(r"\D", "", raw)
        # Filter out year-like 4-digit hits and very short fragments.
        if 7 <= len(digits) <= 15:
            return raw.strip()
    return None


def _looks_like_name(line: str) -> bool:
    """Heuristic: 2–5 words, alphabetic, mostly TitleCased, no digits."""
    line = line.strip(" \t-•·")
    if not line or len(line) > 60:
        return False
    if any(c.isdigit() for c in line):
        return False
    if EMAIL_RE.search(line) or URL_RE.search(line):
        return False
    words = [w for w in re.split(r"\s+", line) if w]
    if not (1 < len(words) <= 5):
        return False
    titled = sum(1 for w in words if w[:1].isupper() and w[1:].islower())
    upper = sum(1 for w in words if w.isupper())
    return (titled + upper) >= max(2, len(words) - 1)


def _extract_name(lines: Sequence[str], *, email: Optional[str]) -> Optional[str]:
    # 1. Look at the first 8 lines for an obviously-name-shaped line.
    for line in lines[:8]:
        if _looks_like_name(line):
            return line.strip(" \t-•·")
    # 2. Derive from the email local part if it has a separator
    #    (e.g. "ahmed.al-hassan@…" → "Ahmed Al Hassan").
    if email:
        local = email.split("@", 1)[0]
        cleaned = re.sub(r"[._\-+]+", " ", local)
        cleaned = re.sub(r"\d+", "", cleaned).strip()
        if cleaned and " " in cleaned:
            return " ".join(w.capitalize() for w in cleaned.split())
    return None


def _extract_one_of(text: str, vocab: Iterable[str]) -> Optional[str]:
    lower = text.lower()
    for term in vocab:
        if re.search(rf"\b{re.escape(term.lower())}\b", lower):
            return term
    return None


def _extract_location(lines: Sequence[str]) -> Optional[str]:
    # Try "Address: …" / "Location: …" first.
    for line in lines[:30]:
        m = re.match(
            r"\s*(?:address|location|residing\s*in|based\s*in)\s*[:\-]\s*(.+)",
            line,
            re.IGNORECASE,
        )
        if m:
            value = m.group(1).strip(" .,-")
            if value:
                return value[:255]

    # Fall back to a line that contains a known city/country.
    for line in lines[:40]:
        for city in GULF_CITIES:
            if re.search(rf"\b{re.escape(city)}\b", line, re.IGNORECASE):
                return line.strip(" .,-")[:255]
    return None


def _extract_years(pattern: re.Pattern[str], text: str) -> Optional[float]:
    """Largest number-of-years hit in the document."""
    best: Optional[float] = None
    for raw in pattern.findall(text):
        try:
            value = float(raw)
        except (TypeError, ValueError):
            continue
        if 0 < value <= 60 and (best is None or value > best):
            best = value
    return best


def _extract_regional_years(text: str, *, region_terms: Sequence[str]) -> Optional[float]:
    """Look for '<N> years … <region>' or '<region> … <N> years'."""
    best: Optional[float] = None
    for term in region_terms:
        for m in re.finditer(
            rf"(\d{{1,2}}(?:\.\d)?)\s*\+?\s*(?:years?|yrs?)\b[^\n]{{0,80}}\b{re.escape(term)}\b",
            text,
            re.IGNORECASE,
        ):
            try:
                v = float(m.group(1))
            except ValueError:
                continue
            if 0 < v <= 60 and (best is None or v > best):
                best = v
        for m in re.finditer(
            rf"\b{re.escape(term)}\b[^\n]{{0,80}}(\d{{1,2}}(?:\.\d)?)\s*\+?\s*(?:years?|yrs?)\b",
            text,
            re.IGNORECASE,
        ):
            try:
                v = float(m.group(1))
            except ValueError:
                continue
            if 0 < v <= 60 and (best is None or v > best):
                best = v
    return best


def _extract_salary(text: str) -> Optional[int]:
    m = SALARY_RE.search(text)
    if not m:
        return None
    raw = re.sub(r"[\s,]", "", m.group(1))
    try:
        value = int(raw)
    except ValueError:
        return None
    # Sanity bounds: monthly salary in QAR/AED is roughly 1,000–100,000.
    if 500 <= value <= 1_000_000:
        return value
    return None


def _extract_notice(text: str) -> Optional[str]:
    m = NOTICE_RE.search(text)
    if not m:
        return None
    return m.group(1).strip().capitalize()[:120]


def _extract_visa(text: str) -> Optional[str]:
    m = VISA_RE.search(text)
    if not m:
        return None
    snippet = m.group(0).strip()
    snippet = re.sub(r"\s+", " ", snippet)
    return snippet[:120]


def _split_sections(lines: Sequence[str]) -> Dict[str, List[str]]:
    """Group lines under their nearest section header.

    Returns a dict keyed by canonical section name (``experience``,
    ``education``, etc.). Lines before any recognised header land under
    ``"preamble"``.
    """
    sections: Dict[str, List[str]] = {"preamble": []}
    current = "preamble"
    for raw in lines:
        line = raw.strip().rstrip(":")
        normalised = line.lower()
        matched_header: Optional[str] = None
        for key, headers in SECTION_HEADERS.items():
            if normalised in headers:
                matched_header = key
                break
        if matched_header is not None:
            current = matched_header
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(raw)
    return sections


def _extract_skills(text: str, section_lines: Optional[List[str]]) -> List[str]:
    found: List[str] = []
    seen: set[str] = set()

    def _add(name: str) -> None:
        key = name.strip().lower()
        if key and key not in seen:
            seen.add(key)
            found.append(name.strip())

    if section_lines:
        # Split tokens on bullets / commas / pipes / slashes.
        for line in section_lines:
            for tok in re.split(r"[•·,/|;]+", line):
                tok = tok.strip(" \t-")
                if tok and 2 <= len(tok) <= 60:
                    _add(tok)

    # Dictionary pass — works even when there's no explicit "Skills" section.
    lower = text.lower()
    for kw in SKILLS_DICT:
        if re.search(rf"\b{re.escape(kw.strip().lower())}\b", lower):
            _add(kw.strip())

    return found[:50]


def _extract_languages(text: str, section_lines: Optional[List[str]]) -> List[str]:
    found: List[str] = []
    if section_lines:
        joined = " ".join(section_lines)
        for lang in LANGUAGE_DICT:
            if re.search(rf"\b{re.escape(lang)}\b", joined, re.IGNORECASE):
                if lang not in found:
                    found.append(lang)
    if not found:
        for lang in LANGUAGE_DICT:
            if re.search(rf"\b{re.escape(lang)}\b", text, re.IGNORECASE):
                if lang not in found:
                    found.append(lang)
    return found[:20]


def _extract_certifications(
    text: str, section_lines: Optional[List[str]]
) -> List[str]:
    found: List[str] = []
    seen: set[str] = set()

    def _add(item: str) -> None:
        key = item.strip().lower()
        if key and key not in seen and 2 <= len(item) <= 200:
            seen.add(key)
            found.append(item.strip())

    if section_lines:
        for line in section_lines:
            line = line.strip(" \t•·-")
            if line:
                _add(line)

    for token in CERT_TOKENS:
        # Use re to match either the token alone or as part of a longer phrase.
        match = re.search(
            rf"[^\n]*\b{re.escape(token)}\b[^\n]*", text, re.IGNORECASE
        )
        if match:
            _add(match.group(0).strip())

    return found[:30]


def _extract_education(section_lines: Sequence[str]) -> List[ParsedEducation]:
    if not section_lines:
        return []
    items: List[ParsedEducation] = []
    seen_raw: set[str] = set()
    for raw in section_lines:
        raw = raw.strip(" \t•·-")
        if not raw or raw.lower() in seen_raw:
            continue
        if not any(tok.lower() in raw.lower() for tok in DEGREE_TOKENS):
            continue
        seen_raw.add(raw.lower())
        degree = next(
            (tok for tok in DEGREE_TOKENS if tok.lower() in raw.lower()),
            None,
        )
        year_match = re.search(r"\b(19|20)\d{2}\b", raw)
        institution: Optional[str] = None
        # Common shape: "Bachelor of Engineering, ABC University, 2018".
        parts = [p.strip() for p in re.split(r"[,•|]", raw) if p.strip()]
        if len(parts) >= 2:
            for part in parts:
                if re.search(r"university|college|institute|school", part, re.IGNORECASE):
                    institution = part
                    break
        items.append(
            ParsedEducation(
                raw=raw[:255],
                degree=degree,
                institution=institution[:200] if institution else None,
                year=int(year_match.group(0)) if year_match else None,
            )
        )
    return items[:10]


_COMPANY_HINT = re.compile(
    r"\b(?:at|@|with)\s+[A-Z][\w&.,'\- ]{2,60}",
)
_DATE_RANGE = re.compile(
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec|"
    r"January|February|March|April|June|July|August|September|"
    r"October|November|December)?\s*'?\d{2,4}\s*[\-–to]+\s*"
    r"(?:Present|Current|"
    r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec|"
    r"January|February|March|April|June|July|August|September|"
    r"October|November|December)?\s*'?\d{2,4}))",
    re.IGNORECASE,
)


def _extract_companies(section_lines: Sequence[str]) -> List[ParsedCompany]:
    if not section_lines:
        return []
    items: List[ParsedCompany] = []
    pending: Dict[str, Optional[str]] = {"name": None, "title": None, "duration": None}

    def _flush() -> None:
        if pending["name"]:
            items.append(
                ParsedCompany(
                    name=pending["name"][:255],
                    title=pending["title"][:255] if pending["title"] else None,
                    duration=pending["duration"][:120] if pending["duration"] else None,
                )
            )
        pending["name"] = None
        pending["title"] = None
        pending["duration"] = None

    for raw in section_lines:
        line = raw.strip(" \t•·-")
        if not line:
            continue

        date_match = _DATE_RANGE.search(line)
        company_hint = _COMPANY_HINT.search(line)

        # Heuristic: a "TitleCase Company  |  Title" or "Title — Company" line
        # often contains a comma or pipe.
        candidate_parts = [p.strip() for p in re.split(r"[|—–-]| at | @ ", line) if p.strip()]

        if date_match:
            pending["duration"] = date_match.group(0)
            # If this is the FIRST hit and the previous one had a name,
            # flush it before starting a new record.
            line_without_dates = _DATE_RANGE.sub("", line).strip(" ,;.-|—")
            if line_without_dates:
                if pending["name"] is None and len(candidate_parts) >= 2:
                    pending["title"] = candidate_parts[0][:255]
                    pending["name"] = candidate_parts[1][:255]
                elif pending["name"] is None:
                    pending["name"] = line_without_dates[:255]
            _flush()
            continue

        if company_hint and pending["name"] is None:
            name_part = company_hint.group(0).split(maxsplit=1)[1]
            pending["name"] = name_part[:255]
            title_part = line[: company_hint.start()].strip(" ,;:-|—")
            if title_part:
                pending["title"] = title_part[:255]
            continue

        # "Title at Company" caught above; otherwise treat first plain line
        # after a blank as the company.
        if pending["name"] is None and len(candidate_parts) >= 2 and any(
            w in candidate_parts[1].lower()
            for w in ("ltd", "llc", "wll", "group", "company", "corporation", "co.", "inc", "limited")
        ):
            pending["title"] = candidate_parts[0][:255]
            pending["name"] = candidate_parts[1][:255]
            continue

    _flush()
    return items[:20]
